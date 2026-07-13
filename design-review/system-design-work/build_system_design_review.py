from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile
import hashlib

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/jacob/Documents/App Monitor")
REFERENCE = Path("/Users/jacob/.codex/plugins/cache/openai-curated-remote/openai-templates/0.1.0/skills/artifact-template-system-design/assets/reference.docx")
WORK = ROOT / "design-review/system-design-work"
FINAL = ROOT / "design-review/App Monitor System and Product Design Review.docx"
DIAGRAM = WORK / "app-monitor-proposed-architecture.png"

EXPECTED_SHA = "13504f6c221a42c1726460a9e865e563355539ff97d702d6c9b2267b4b261d76"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def set_run_font(run, name="Helvetica Neue", size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def replace_paragraph(paragraph, text):
    runs = list(paragraph.runs)
    source_rpr = deepcopy(runs[0]._r.rPr) if runs and runs[0]._r.rPr is not None else None
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if source_rpr is not None:
        run._r.insert(0, source_rpr)
    elif text:
        set_run_font(run)


def set_cell(cell, text, *, bold=False, center=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.runs:
        paragraph.runs[0].bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_numbered_paragraph(doc, text):
    p = doc.add_paragraph(style="normal")
    p.style = doc.styles["normal"]
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="normal")
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(7)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5)
    return p


def add_bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="normal")
    p.paragraph_format.left_indent = Inches(0.24 + 0.18 * level)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("•  " + text)
    set_run_font(r, size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style="Heading 1" if level == 1 else "Heading 2")
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.keep_with_next = True
    replace_paragraph(p, text)
    return p


def page_break(doc):
    doc.add_page_break()


def draw_architecture(path: Path):
    width, height = 1600, 760
    img = Image.new("RGB", (width, height), "#F3F7FA")
    d = ImageDraw.Draw(img)
    font_paths = [
        "/System/Library/Fonts/Helvetica.ttc",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    font_path = next(p for p in font_paths if Path(p).exists())
    title = ImageFont.truetype(font_path, 34)
    body = ImageFont.truetype(font_path, 25)
    small = ImageFont.truetype(font_path, 19)
    d.rectangle((0, 0, width, 108), fill="#0B2E50")
    d.text((58, 31), "App Monitor - proposed modular architecture", font=title, fill="white")
    d.text((58, 74), "Local-first orchestration with explicit policy, durable operation records, and contextual UI state", font=small, fill="#D7E6F4")

    boxes = {
        "ui": (55, 200, 340, 355, "SwiftUI shell", "Health / Updates\nStorage / Activity"),
        "coord": (470, 175, 825, 380, "Domain coordinators", "Inventory & Usage\nStorage & Safety\nUpdate Orchestrator"),
        "svc": (955, 175, 1295, 380, "Services / adapters", "Scanners • providers\ncleanup • uninstall\nself-update"),
        "ext": (1335, 200, 1550, 355, "macOS + tools", "Workspace\nmas • brew\nsoftwareupdate"),
        "policy": (470, 500, 825, 650, "Policy & settings", "Safety rules • cadence • permissions\nretention • selected UI direction"),
        "store": (955, 500, 1295, 650, "SQLite + operation ledger", "Snapshots • run/item results\nchangelogs • restore history"),
    }

    def box(coords, heading, sub, fill="#FFFFFF"):
        x1, y1, x2, y2 = coords
        d.rounded_rectangle(coords, radius=16, fill=fill, outline="#2E6F9E", width=3)
        d.text((x1 + 22, y1 + 22), heading, font=body, fill="#0B2E50")
        y = y1 + 70
        for line in sub.split("\n"):
            d.text((x1 + 22, y), line, font=small, fill="#465D73")
            y += 30

    for key, values in boxes.items():
        x1, y1, x2, y2, h, s = values
        box((x1, y1, x2, y2), h, s, "#E5EFF7" if key in {"policy", "store"} else "#FFFFFF")

    def arrow(a, b):
        d.line((a, b), fill="#2E6F9E", width=7)
        x, y = b
        if a[0] < b[0]:
            points = [(x, y), (x - 25, y - 16), (x - 25, y + 16)]
        elif a[1] < b[1]:
            points = [(x, y), (x - 16, y - 25), (x + 16, y - 25)]
        else:
            points = [(x, y), (x - 16, y + 25), (x + 16, y + 25)]
        d.polygon(points, fill="#2E6F9E")

    arrow((340, 278), (470, 278))
    arrow((825, 278), (955, 278))
    arrow((1295, 278), (1335, 278))
    arrow((650, 380), (650, 500))
    arrow((1125, 380), (1125, 500))
    arrow((955, 575), (825, 575))
    d.text((58, 704), "Invariant: no destructive side effect without a reviewable target, a captured policy snapshot, and a durable result record.", font=small, fill="#465D73")
    img.save(path)


def add_picture(doc, path: Path, width=7.0, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    alt_text = caption or path.stem.replace("-", " ").replace("_", " ")
    shape._inline.docPr.set("title", alt_text)
    shape._inline.docPr.set("descr", alt_text)
    if caption:
        c = doc.add_paragraph(style="normal")
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        set_run_font(r, size=9, color=(75, 95, 115))
    return p


def clone_table_after(doc, template_index, headers, rows):
    source = doc.tables[template_index]
    new_tbl = deepcopy(source._tbl)
    doc._body._body.append(new_tbl)
    table = doc.tables[-1]
    desired = 1 + len(rows)
    while len(table.rows) < desired:
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    while len(table.rows) > desired:
        table._tbl.remove(table.rows[-1]._tr)
    for ci, header in enumerate(headers):
        set_cell(table.rows[0].cells[ci], header, bold=True)
    for ri, row in enumerate(rows, start=1):
        for ci, value in enumerate(row):
            set_cell(table.rows[ri].cells[ci], value)
            shade_cell(table.rows[ri].cells[ci], "E6F0F7" if ri % 2 else "F3F7FA")
    return table


assert sha256(REFERENCE) == EXPECTED_SHA, "Retained reference changed; redistill before authoring."
draw_architecture(DIAGRAM)
doc = Document(REFERENCE)

# Cover
replace_paragraph(doc.paragraphs[8], "App Monitor")
replace_paragraph(doc.paragraphs[9], "System & Product Review")

meta = doc.tables[0]
for cell, value in zip((meta.cell(0, 0), meta.cell(0, 2), meta.cell(0, 4)), ("Implemented", "Jacob Crandall", "July 11, 2026")):
    target = cell.paragraphs[-1]
    replace_paragraph(target, value)

info = doc.tables[1]
set_cell(info.cell(0, 1), "OpenAI Codex")
set_cell(info.cell(1, 1), "Product owner; macOS engineering; security/privacy review")
set_cell(info.cell(2, 1), "README.md; PRIVACY.md; RELEASING.md; design handoffs")
set_cell(info.cell(3, 1), "Architecture, live macOS UI, safety, operations, and five redesign directions.")

# 1. Abstract
replace_paragraph(doc.paragraphs[22], "App Monitor is a local-first macOS utility that combines app inventory, foreground usage, storage analysis, update management, reversible cleanup, uninstall planning, and action history. The core is technically viable and currently healthy. This review also shipped a focused UX pass: the inspector is now contextual, primary destinations have unique names, provider filters are collapsed, Settings is pinned, toolbar options are consolidated, and raw login-item state is translated into plain language. The remaining risk is accumulated product and orchestration complexity concentrated in one SwiftUI screen and one main-actor model.")
replace_paragraph(doc.paragraphs[23], "The next architectural step is an action-first shell backed by domain coordinators, a versioned operation ledger, explicit scan/update coverage, and route-scoped detail state. The design remains native, single-user, and local-only on macOS 14+. It does not add accounts, a hosted backend, telemetry, or autonomous destructive behavior. External dependencies remain the local file system, macOS frameworks, Homebrew, mas, softwareupdate, and app-provided release feeds.")

# 2. Goals and non-goals
goals = [
    "Make the next useful action obvious within five seconds of launch.",
    "Keep cleanup, uninstall, and update actions previewable, reversible where possible, and fully auditable.",
    "Keep cached navigation responsive while scans and provider work run off the main actor with cancellation and coverage reporting.",
    "Ship the redesign incrementally behind a feature flag while preserving the existing SQLite database and provider test suite.",
]
nongoals = [
    "Replace dedicated package managers, malware scanners, or whole-disk visualizers.",
    "Add cloud sync, user accounts, cross-device history, or product telemetry.",
    "Auto-delete user data or install manual/restart updates without explicit policy and review.",
    "Redesign every specialist analytics view before the navigation and orchestration foundation is stable.",
]
for i in range(4):
    set_cell(doc.tables[2].cell(i + 1, 0), goals[i])
    set_cell(doc.tables[2].cell(i + 1, 1), nongoals[i])

# 3. Background
replace_paragraph(doc.paragraphs[28], "The audit started with a capable but crowded three-pane workspace. The sidebar mixed top-level destinations, quick filters, update provider sources, analytics, and maintenance; the persistent right inspector showed stale app detail on Updates, History, Settings, and analytics. The implemented pass now defaults every route to a full-width workspace, opens details only when relevant, gives destinations unique names, collapses source filters, pins Settings, and organizes settings copy into understandable sections. The strongest flow remains Quarantine Review because it explains the item, exact path, safety steps, and reversible queue in one place.")
replace_paragraph(doc.paragraphs[29], "The implementation reflects the same breadth: DashboardView.swift is 9,430 lines, AppModel.swift is 4,449 lines with more than 50 published properties, and AppDataStore.swift is 2,663 lines with many schema responsibilities. Caching already removed redraw-time SQLite work, but AppModel remains the owner of inventory, analytics, scanning, cleanup, updates, self-update, uninstall, scheduling, settings, selection, and history. The proposed boundary is: SwiftUI shell and route view models -> domain coordinators -> scanner/provider/executor adapters -> SQLite operation/state store and external macOS tools. The invariant is fail closed for destructive side effects and keep the UI driven by durable or explicitly transient state.")

# 4. Architecture image
picture_paragraph = doc.paragraphs[32]
replace_paragraph(picture_paragraph, "")
picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
architecture_shape = picture_paragraph.add_run().add_picture(str(DIAGRAM), width=Inches(7.05))
architecture_shape._inline.docPr.set("title", "Proposed modular architecture for App Monitor")
architecture_shape._inline.docPr.set("descr", "SwiftUI shell routes work through domain coordinators and service adapters to macOS tools, with policy settings and a SQLite operation ledger below.")
replace_paragraph(doc.paragraphs[33], "Figure 1. Proposed modular architecture for App Monitor.")

components = [
    ("App shell + route view models", "Own navigation, route-specific state, contextual selection, accessibility labels, and rendering only.", "Published snapshots and transient selection", "Show last valid snapshot and a route-scoped error; never block the whole shell."),
    ("Inventory & Usage coordinator", "Inventory scans, tracking, timeline/summary queries, retention, exclusions, and cached usage snapshots.", "apps, usage_segments, imported usage", "Tracking can pause; queries fall back to last durable data with freshness shown."),
    ("Storage & Safety coordinator", "Storage jobs, coverage, health findings, cleanup policy, quarantine, restore, and uninstall plans.", "storage_items, findings, suggestions, action history", "Permission gaps are first-class coverage states; destructive execution fails closed."),
    ("Update orchestrator", "Normalize providers, compute eligibility, run selected updates, reconcile results, and capture changelogs.", "app_updates, update runs/results, changelogs", "Provider failures remain isolated and auditable; manual/restart items are never auto-selected."),
    ("Persistence + operation ledger", "Versioned migrations, transactions, run/item lifecycle, settings snapshots, integrity checks, and export boundaries.", "SQLite WAL database", "A required write failure blocks side effects; recovery uses idempotent rerun and reconciliation."),
]
for ri, row in enumerate(components, start=1):
    for ci, value in enumerate(row):
        set_cell(doc.tables[3].cell(ri, ci), value)

# 5. Request lifecycle
lifecycle = [
    "A manual action, scheduled timer, app activation event, or usage event enters a domain coordinator with route filters, target IDs, and current settings.",
    "The coordinator normalizes paths/provider records, checks permissions and policy, rejects protected/symlinked/unsupported targets, and resolves the user-visible action type.",
    "It loads the last durable snapshot plus a frozen policy/settings snapshot; the UI immediately publishes a lightweight queued/running state.",
    "Scanner and provider work runs off the main actor with bounded concurrency, progress, cancellation, and provider-specific timeouts.",
    "Before cleanup, uninstall, restore, or update side effects begin, a run record and selected item set are persisted; inability to write fails closed.",
    "Each downstream result is recorded independently. Partial provider failure does not erase successful results, and restart/administrator requirements remain explicit terminal states.",
    "The coordinator reconciles installed/file-system state, closes the run, refreshes cached view snapshots, and exposes completion, failures, and restore/retry actions in History.",
]
for index, text in zip(range(40, 47), lifecycle):
    replace_paragraph(doc.paragraphs[index], text)

# 6. Contract
contract_rows = [
    ("operation_id", "UUID text", "Yes", "Stable run identifier for scan, update, cleanup, uninstall, restore, or import."),
    ("kind", "Enum", "Yes", "Defines policy, lifecycle, eligible targets, and valid result states."),
    ("target_id", "Text", "No", "App, provider, path, or aggregate scope; paths are standardized before persistence."),
    ("requested_at", "Timestamp", "Yes", "Local audit time; completed_at is stored when the run reaches a terminal state."),
    ("policy_snapshot", "Versioned JSON", "Yes", "Frozen cadence, safety, permission, and auto-eligibility settings used by this run."),
    ("status", "Enum", "Yes", "queued, running, completed, partial, failed, cancelled, or needsRestart."),
    ("item_results", "Nested records", "Yes", "Per-target status, message, bytes/version transition, side-effect evidence, and recovery link."),
]
for ri, row in enumerate(contract_rows, start=1):
    for ci, value in enumerate(row):
        set_cell(doc.tables[4].cell(ri, ci), value)

guarantees = [
    "A side-effecting action must have a validated target set and durable run record before execution.",
    "Run IDs and item IDs are stable; reconciliation updates records instead of creating ambiguous duplicate history.",
    "Provider/source, from/to versions, target path, policy version, timestamps, and terminal outcome are captured for audit and recovery.",
    "SQLite is the source of truth for App Monitor observations and operations, not for external package-manager or file-system truth; reconciliation remains required.",
]
for index, text in zip(range(54, 58), guarantees):
    replace_paragraph(doc.paragraphs[index], text)
replace_paragraph(doc.paragraphs[58], "The schema is implemented in AppDataStore migrations and should be formalized as versioned migration files with a schema_version table.")
replace_paragraph(doc.paragraphs[59], "Keep migration fixtures and contract tests alongside AppMonitorCoreTests; publish human-readable field semantics in docs/.")

# 7. Consistency
replace_paragraph(doc.paragraphs[63], "App Monitor needs strong local ordering around side effects, but not distributed-system semantics. A serial datastore queue and SQLite WAL provide ordered writes; coordinators freeze policy per run and publish UI snapshots only after coherent state transitions. Read-only scans and provider checks may be repeated. Cleanup, uninstall, restore, and install actions require stable IDs, preflight validation, durable attempt records, and post-action reconciliation. Partial success is acceptable only when each item has an explicit terminal result and the UI does not imply atomic completion.")
consistency_rows = [
    ("Repeated scan or update check", "Coalesce while active or create a new run after the prior run closes; replace current snapshots deterministically.", "Read-only work is safe to repeat and durable history remains unambiguous."),
    ("Run-record write fails", "Do not start destructive/update execution; surface a retryable storage error.", "Failing closed prevents unaudited side effects."),
    ("Provider or item partially fails", "Persist per-item outcomes, mark the run partial, reconcile successful items, and offer scoped retry.", "A truthful partial state is safer than rollback claims the external tool cannot guarantee."),
    ("Settings change mid-run", "Continue with the captured policy snapshot; new settings apply to the next run.", "Stable behavior and audit reconstruction require immutable per-run policy."),
]
for ri, row in enumerate(consistency_rows, start=1):
    for ci, value in enumerate(row):
        set_cell(doc.tables[5].cell(ri, ci), value)

# 8. Security and privacy
security = [
    "Single-user local trust boundary. Use standard macOS authorization only at the action that needs it; protected paths and Apple/system apps remain blocked.",
    "Usage history, paths, app inventory, and action history are sensitive. Keep them local, redact secrets from messages, and never add telemetry without a separate explicit design review.",
    "Credentials belong in Keychain or ephemeral authorization context only, never SQLite, UserDefaults, process arguments, exports, or logs. Self-update packages require checksum/signature validation.",
    "Cleanup defaults to quarantine; uninstall defaults high-risk paths off; symlink destinations are not followed; direct-download actions allow only trusted URL schemes and explicit user review.",
    "Add pause tracking, per-app exclusion, clear-history, and 30/90/365-day/forever retention. Show scan coverage and permission gaps so totals are not presented as complete when access is partial.",
]
for index, text in zip(range(67, 72), security):
    replace_paragraph(doc.paragraphs[index], text)

# 9. Operational readiness
readiness = [
    ("Navigation response", "Target p95 <100 ms from click to cached content; no SQLite/file scan in SwiftUI body.", "App Monitor", "Required"),
    ("Progress visibility", "Visible progress/freshness within 250 ms; cancellation acknowledged within 1 s.", "App Monitor", "Required"),
    ("Provider completion", ">=95% successful provider checks excluding unavailable/uninstalled tools; failures isolated by source.", "Update orchestrator", "Required"),
    ("Audit completeness", "100% of side-effecting item attempts have run ID, target, policy snapshot, and terminal result.", "Operation ledger", "Required"),
    ("Restore/integrity", "100% deterministic fixture restores; zero silent migration or relationship loss in upgrade tests.", "Storage & persistence", "Required"),
]
for ri, row in enumerate(readiness, start=1):
    for ci, value in enumerate(row):
        set_cell(doc.tables[6].cell(ri, ci), value)
set_cell(doc.tables[6].cell(6, 0), "Rollout: feature-flag the new shell; seed deterministic data; run current/new shells against the same database; preserve a one-release rollback path; promote only after accessibility and destructive-action audit gates pass.")

# 10. Alternatives
alternatives = [
    ("Keep extending the current shell", "Lowest short-term implementation cost.", "Perpetuates context mismatch and large AppModel/DashboardView blast radius."),
    ("Adopt a full architecture framework", "Strong reducer/state conventions and testability.", "Large migration cost; domain coordinators deliver most value without framework lock-in."),
    ("Split into separate usage, cleanup, and updater apps", "Each product becomes easier to explain.", "Loses the core differentiator: usage, storage, risk, and update context together."),
    ("Add a cloud backend", "Cross-device sync and remote analytics become possible.", "Conflicts with local-first privacy, adds account/security operations, and is unnecessary for the current goal."),
]
for ri, row in enumerate(alternatives, start=1):
    for ci, value in enumerate(row):
        set_cell(doc.tables[7].cell(ri, ci), value)

# 11. Open questions
questions = [
    "Which of the five UI directions should define the new shell and primary navigation model?",
    "Should route-specific inspector visibility persist independently after the contextual default has proven stable?",
    "What is the default usage-history retention period, and should excluded apps be dropped at capture time or filtered later?",
    "Which update classes may ever run automatically: managed Homebrew only, App Store, direct-download, or none without confirmation?",
]
for index, text in zip(range(81, 85), questions):
    replace_paragraph(doc.paragraphs[index], text)

# 12. Decision
replace_paragraph(doc.paragraphs[87], "Ship the implemented contextual-inspector, simplified-navigation, toolbar, and settings-language pass as the current baseline. Then adopt an action-first product shell and refactor behind it rather than rewriting the core. Extract Inventory & Usage, Storage & Safety, Updates, and Persistence/Operations coordinators from AppModel; decompose DashboardView into route files; add a versioned operation ledger, cancellation, permission/coverage states, and deterministic preview data; then roll the selected shell direction out behind a feature flag while the current packaged app remains the rollback path.")
milestones = [
    ("M1", "Extract route files and domain coordinators; preserve behavior and current database.", "56 tests and scripts/ci pass; no route performs datastore/file work in body; cached navigation target met."),
    ("M2", "Versioned migrations, operation ledger, scan coverage, cancellation, retention/exclusion settings.", "Failure/replay fixtures pass; destructive attempts are 100% auditable; permission gaps are visible."),
    ("M3", "Implement the selected UI direction behind a feature flag with deterministic demo state.", "Screenshot QA, keyboard traversal, VoiceOver labels, contrast, narrow-window, light/dark tests pass."),
    ("M4", "Refine specialist Update, Storage, Activity, and History workspaces and retire the legacy shell.", "Stable beta period; no P0/P1 safety regressions; rollback tested; product metrics accepted."),
]
for ri, row in enumerate(milestones, start=1):
    for ci, value in enumerate(row):
        set_cell(doc.tables[8].cell(ri, ci), value)

# Footer
for section in doc.sections:
    for paragraph in section.footer.paragraphs:
        if paragraph.text.strip():
            replace_paragraph(paragraph, "App Monitor | System & Product Design Review")

# Appendix A
page_break(doc)
add_heading(doc, "Appendix A. Product Design Audit", 1)
add_body(doc, "Overall verdict: the product is capable and unusually transparent for a Mac maintenance utility. The completed UX pass removes the most visible orientation and space problems without changing the product's visual system or safety model. The best existing pattern remains Quarantine Review because it connects evidence, safety, and action.")
add_heading(doc, "Highest-impact changes", 2)
audit_items = [
    "IMPLEMENTED — Give primary destinations unique names, collapse provider/source filters, and keep Settings pinned below the scrolling navigation list.",
    "IMPLEMENTED — Make the inspector contextual and hidden on navigation; expose a labeled Details control only when the route has inspectable content.",
    "IMPLEMENTED — Consolidate filtering, sorting, saved filters, and export under Options so Scan, Updates, and Cleanup remain readable.",
    "IMPLEMENTED — Group Settings around General, Full Scan Schedule, Update App Monitor, and Update Installed Software; translate login-item status into plain language.",
    "P1 — Turn the raw warning total into an actionable summary by severity, affected app, confidence, and recommended next action. Do not use a raw count as the primary signal.",
    "P1 — Keep long-running work in a compact status center. Progress must not push every route downward or disable unrelated inspection.",
    "P1 — Use progressive disclosure for changelogs and history. Keep the current task visible and open detailed timelines only on selection.",
    "P1 — Audit small labels, pale gray text, icon-only controls, focus order, keyboard hit targets, and non-color state labels in both light and dark themes.",
    "P1 — Say “space available for review,” not guaranteed savings. Show freshness and permission coverage next to totals.",
    "P2 — Scope search and filters explicitly so users know whether they affect apps, providers, storage paths, or usage analytics.",
]
for item in audit_items:
    add_bullet(doc, item)
add_heading(doc, "Evidence and limits", 2)
add_body(doc, "Evidence comes from a live July 11, 2026 capture of Overview, Installed App Updates, Quarantine Review, History, Settings, and Usage Trends in the locally built and signed app bundle. Screenshots support layout, hierarchy, wording, density, and visible state findings. They do not prove VoiceOver traversal, measured contrast, destructive-action correctness, permissions behavior, or narrow-window behavior; those require instrumented testing.")

audit_screens = [
    (ROOT / "design-review/live-audit-2026-07-11/07-overview-after.png", "Implemented Overview — full-width workspace, readable primary toolbar, unique navigation labels, collapsed advanced sources, and pinned Settings."),
    (ROOT / "design-review/live-audit-2026-07-11/09-quarantine-after.png", "Implemented Quarantine Review — the queue uses the full canvas and safety detail remains one explicit Details action away."),
    (ROOT / "design-review/live-audit-2026-07-11/10-settings-after.png", "Implemented Settings — route-relevant full-width content, plain-language sections, and a human-readable launch-at-login status."),
]
for path, caption in audit_screens:
    page_break(doc)
    add_picture(doc, path, width=7.0, caption=caption)

# Appendix B
page_break(doc)
add_heading(doc, "Appendix B. UI Revision Directions", 1)
add_body(doc, "All five directions are grounded in the live product and preserve local-first, safety-first behavior. The choice should be based on the product’s primary promise, not surface styling alone.")
directions = [
    ("Direction 1 — Operations Console", "Best for a product promise centered on “what needs attention now.” Strongest synthesis of updates, warnings, and safe maintenance; requires careful prioritization rules.", ROOT / "design-review/ui-directions/01-operations-console.png"),
    ("Direction 2 — Calm Mac Health", "Best for mainstream users who want reassurance and guided recommendations. Lowest cognitive load; risks hiding expert detail too early.", ROOT / "design-review/ui-directions/02-calm-mac-health.png"),
    ("Direction 3 — Update Command Center", "Best if broad update management becomes the lead differentiator. Excellent provider/release-note clarity; de-emphasizes usage and storage context.", ROOT / "design-review/ui-directions/03-update-command-center.png"),
    ("Direction 4 — Activity Timeline Workspace", "Best if personal usage intelligence becomes the lead product. Strong coordinated analysis; requires first-class retention/privacy controls.", ROOT / "design-review/ui-directions/04-activity-timeline-workspace.png"),
    ("Direction 5 — Compact Pro Utility", "Best for expert users who value density, keyboard speed, and Finder-like inspection. Scales well; has the highest accessibility and onboarding burden.", ROOT / "design-review/ui-directions/05-compact-pro-utility.png"),
]
for i, (title, text, path) in enumerate(directions):
    if i:
        page_break(doc)
    direction_heading = add_heading(doc, title, 2)
    direction_heading.style = doc.styles["normal"]
    direction_heading.paragraph_format.left_indent = Inches(0)
    direction_heading.paragraph_format.right_indent = Inches(0)
    direction_heading.paragraph_format.first_line_indent = Inches(0)
    direction_heading.paragraph_format.space_after = Pt(8)
    direction_heading.paragraph_format.keep_with_next = True
    direction_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in direction_heading.runs:
        set_run_font(run, size=13, bold=True, color=(89, 113, 134))
    add_body(doc, text)
    add_picture(doc, path, width=7.0)

# Appendix C
page_break(doc)
add_heading(doc, "Appendix C. Prioritized Implementation Backlog", 1)
add_body(doc, "The backlog separates product clarity from deep infrastructure work so the team can improve the experience without compromising safety or creating a rewrite-sized release.")
backlog = [
    ("Done", "Contextual inspector + simplified navigation", "Route relevance, substantially more usable width, unique labels, and pinned Settings."),
    ("P0 / M", "Split DashboardView and AppModel into route/domain files", "Files become reviewable; route changes no longer touch unrelated workflows."),
    ("P0 / M", "Versioned operation ledger and fail-closed side effects", "100% destructive/update attempts reconstructable by run and item."),
    ("P0 / S", "Settings categories + privacy controls", "Pause, exclusions, retention, clear history, and permission coverage are discoverable."),
    ("P1 / M", "Cancelable scan jobs + compact status center", "Progress visible <250 ms; cancellation <1 s; unrelated routes remain usable."),
    ("P1 / S", "Warning severity/confidence/action model", "Raw warning count replaced by prioritized, explainable review groups."),
    ("P1 / M", "Deterministic demo database + screenshot states", "Reliable empty/loading/error/dense/light/dark visual regression coverage."),
    ("P1 / M", "Accessibility validation matrix", "Keyboard, focus, VoiceOver labels, contrast, non-color states, and large text pass."),
    ("P2 / L", "Incremental scans and persisted storage rollups", "Fast category/app refresh; historical deltas without loading all paths."),
]
clone_table_after(doc, 7, ("Priority / effort", "Change", "Success measure"), backlog)

for paragraph in doc.paragraphs:
    if paragraph.text.strip() in {"Core components", "Primary data contract"}:
        paragraph.style = doc.styles["Heading 2"]

FINAL.parent.mkdir(parents=True, exist_ok=True)
doc.save(FINAL)
assert sha256(REFERENCE) == EXPECTED_SHA, "Reference changed during authoring."
print(FINAL)
